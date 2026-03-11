"""
Generador del Brief Operativo para el Gerente de Destinos Express.
Crea un documento Word profesional con buena tipografía y estructura.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── COLORES ──────────────────────────────────────────────────────────────────
NEGRO       = RGBColor(0x0D, 0x0D, 0x0D)
DORADO      = RGBColor(0xFC, 0xA3, 0x11)
DORADO_OSC  = RGBColor(0xE0, 0x8E, 0x00)
BLANCO      = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_CLARO  = RGBColor(0xF5, 0xF5, 0xF5)
GRIS_MED    = RGBColor(0xD0, 0xD0, 0xD0)
GRIS_TEXT   = RGBColor(0x55, 0x55, 0x55)
AZUL_LINK   = RGBColor(0x1A, 0x73, 0xE8)


def hex_to_rgb_str(r, g, b):
    return f"{r:02X}{g:02X}{b:02X}"


NEGRO_HEX      = "0D0D0D"
DORADO_HEX     = "FCA311"
BLANCO_HEX     = "FFFFFF"
GRIS_CLARO_HEX = "F5F5F5"
GRIS_MED_HEX   = "D0D0D0"
GRIS_TEXT_HEX  = "555555"
DORADO_OSC_HEX = "E08E00"
AMARILLO_SUAVE = "FFF8E7"
AZUL_SUAVE     = "EBF3FB"


# ─── HELPERS XML ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val',   'single'))
            el.set(qn('w:sz'),    val.get('sz',    '4'))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run_formatted(para, text, bold=False, italic=False, size=11,
                      color=None, font='Calibri'):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_paragraph(doc, text='', style=None, bold=False, italic=False,
                  size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6, font='Calibri'):
    p = doc.add_paragraph()
    if style:
        try:
            p.style = style
        except Exception:
            pass
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        add_run_formatted(p, text, bold=bold, italic=italic,
                          size=size, color=color, font=font)
    return p


def add_horizontal_line(doc, color_hex=DORADO_HEX, thickness='12'):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pb   = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    thickness)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_page_break(doc):
    doc.add_page_break()


# ─── BLOQUES DE CONTENIDO ─────────────────────────────────────────────────────

def bloque_portada(doc):
    """Portada con fondo negro, título dorado y subtítulo."""

    # Banda negra superior — simulada con tabla de 1 celda
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_bg(cell, NEGRO_HEX)
    cell.width = Cm(16)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(24)
    p1.paragraph_format.space_after  = Pt(4)
    add_run_formatted(p1, 'DESTINOS EXPRESS S.A.S.',
                      bold=True, size=22, color=BLANCO, font='Montserrat')

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    add_run_formatted(p2, 'NIT: 900.982.154-2  ·  Bogotá D.C., Colombia',
                      size=10, color=DORADO, font='Calibri')

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(24)
    add_run_formatted(p3, 'comercial@destinosexpress.com  ·  +57 302 4060101',
                      size=9, color=RGBColor(0xCC, 0xCC, 0xCC), font='Calibri')

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)
    add_run_formatted(p, 'BRIEF OPERATIVO', bold=True, size=28,
                      color=NEGRO, font='Calibri')

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    add_run_formatted(p2, 'Sistema de Cotizaciones — Configuración Inicial',
                      bold=False, size=14, color=GRIS_TEXT, font='Calibri')

    add_horizontal_line(doc, DORADO_HEX, '16')

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Metadatos
    t2 = doc.add_table(rows=3, cols=2)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ['Preparado por:', 'Dirigido a:', 'Fecha:']
    values = ['Equipo técnico — Destinos Express', 'Gerencia General', 'Marzo 2026']
    for i, (lbl, val) in enumerate(zip(labels, values)):
        cl = t2.cell(i, 0)
        cv = t2.cell(i, 1)
        set_cell_bg(cl, GRIS_CLARO_HEX)
        pl = cl.paragraphs[0]
        pl.paragraph_format.space_before = Pt(5)
        pl.paragraph_format.space_after  = Pt(5)
        add_run_formatted(pl, lbl, bold=True, size=10,
                          color=GRIS_TEXT, font='Calibri')
        pv = cv.paragraphs[0]
        pv.paragraph_format.space_before = Pt(5)
        pv.paragraph_format.space_after  = Pt(5)
        add_run_formatted(pv, val, bold=False, size=10,
                          color=NEGRO, font='Calibri')

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Instrucciones generales
    t3 = doc.add_table(rows=1, cols=1)
    cell3 = t3.cell(0, 0)
    set_cell_bg(cell3, AMARILLO_SUAVE)
    set_cell_borders(cell3,
        top={'val': 'single', 'sz': '12', 'color': DORADO_HEX},
        bottom={'val': 'single', 'sz': '4', 'color': DORADO_OSC_HEX},
    )
    pi = cell3.paragraphs[0]
    pi.paragraph_format.space_before = Pt(8)
    pi.paragraph_format.space_after  = Pt(4)
    add_run_formatted(pi, '📋  INSTRUCCIONES DE DILIGENCIAMIENTO',
                      bold=True, size=10, color=DORADO_OSC, font='Calibri')

    instrucciones = [
        '• Lea cada sección completa antes de responder.',
        '• Escriba sus respuestas en los espacios indicados con: ✏️  Respuesta:',
        '• Los campos marcados con ⚠️ son obligatorios para el funcionamiento del sistema.',
        '• Los campos con 💬 son preguntas abiertas — responda con sus propias palabras.',
        '• Si una pregunta no aplica a su operación, escriba "No aplica" en el campo.',
        '• No hay respuestas incorrectas. Buscamos entender cómo opera su empresa.',
    ]
    for ins in instrucciones:
        p = cell3.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_run_formatted(p, ins, size=9.5, color=GRIS_TEXT, font='Calibri')
    cell3.add_paragraph().paragraph_format.space_after = Pt(6)


def titulo_seccion(doc, numero, titulo, descripcion=''):
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    set_cell_bg(cell, NEGRO_HEX)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(10)
    p1.paragraph_format.space_after  = Pt(2)
    add_run_formatted(p1, f'SECCIÓN {numero}  —  {titulo}',
                      bold=True, size=13, color=BLANCO, font='Calibri')
    if descripcion:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(10)
        add_run_formatted(p2, descripcion, size=9, color=DORADO, font='Calibri')
    else:
        cell.paragraphs[0].paragraph_format.space_after = Pt(10)
    add_horizontal_line(doc, DORADO_HEX, '8')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def subtitulo(doc, texto, descripcion=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    add_run_formatted(p, texto, bold=True, size=12, color=DORADO_OSC, font='Calibri')
    if descripcion:
        pd = doc.add_paragraph()
        pd.paragraph_format.space_before = Pt(0)
        pd.paragraph_format.space_after  = Pt(6)
        add_run_formatted(pd, descripcion, italic=True, size=9.5,
                          color=GRIS_TEXT, font='Calibri')


def nota_ejemplo(doc, texto):
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    set_cell_bg(cell, AZUL_SUAVE)
    set_cell_borders(cell,
        left={'val': 'single', 'sz': '16', 'color': '1A73E8'},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    add_run_formatted(p, '💡  Ejemplo:  ', bold=True, size=9,
                      color=RGBColor(0x1A, 0x73, 0xE8), font='Calibri')
    add_run_formatted(p, texto, italic=True, size=9,
                      color=GRIS_TEXT, font='Calibri')
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def campo_respuesta(doc, pregunta, obligatorio=True, es_abierta=False):
    """Genera un campo pregunta + línea de respuesta."""
    # Pregunta
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    icono = '⚠️  ' if obligatorio else '💬  '
    add_run_formatted(p, icono + pregunta, bold=True, size=10,
                      color=NEGRO, font='Calibri')

    # Línea de respuesta
    t = doc.add_table(rows=1 if not es_abierta else 4, cols=1)
    cell = t.cell(0, 0)
    set_cell_bg(cell, GRIS_CLARO_HEX)
    set_cell_borders(cell,
        top={'val': 'single', 'sz': '4', 'color': GRIS_MED_HEX},
        bottom={'val': 'single', 'sz': '4', 'color': GRIS_MED_HEX},
        left={'val': 'single', 'sz': '4', 'color': GRIS_MED_HEX},
        right={'val': 'single', 'sz': '4', 'color': GRIS_MED_HEX},
    )
    pr = cell.paragraphs[0]
    pr.paragraph_format.space_before = Pt(6)
    pr.paragraph_format.space_after  = Pt(6)
    add_run_formatted(pr, '✏️  Respuesta: ', bold=True, size=9.5,
                      color=DORADO_OSC, font='Calibri')
    add_run_formatted(pr, '_' * 70, size=9.5,
                      color=RGBColor(0xCC, 0xCC, 0xCC), font='Calibri')

    if es_abierta:
        for i in range(1, 4):
            row_cell = t.cell(i, 0)
            set_cell_bg(row_cell, GRIS_CLARO_HEX)
            set_cell_borders(row_cell,
                bottom={'val': 'single', 'sz': '4', 'color': GRIS_MED_HEX},
                left={'val': 'single', 'sz': '4', 'color': GRIS_MED_HEX},
                right={'val': 'single', 'sz': '4', 'color': GRIS_MED_HEX},
            )
            pc = row_cell.paragraphs[0]
            pc.paragraph_format.space_before = Pt(4)
            pc.paragraph_format.space_after  = Pt(4)
            add_run_formatted(pc, '_' * 90, size=9,
                              color=RGBColor(0xCC, 0xCC, 0xCC), font='Calibri')

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def tabla_vehiculo(doc, numero, nombre_sugerido='', capacidad_sugerida=''):
    """Tabla de ficha por vehículo."""
    campos = [
        ('⚠️', 'Nombre comercial del vehículo',        nombre_sugerido    or ''),
        ('⚠️', 'Referencia o modelo (marca/línea)',     ''),
        ('⚠️', 'Capacidad de pasajeros',                capacidad_sugerida or ''),
        ('⚠️', 'Tarifa mínima urbana (COP)',            ''),
        ('⚠️', 'Valor por km en ciudad (COP)',          ''),
        ('⚠️', 'Valor por km intermunicipal (COP)',     ''),
        ('⚠️', 'Tarifa por hora (COP)',                 ''),
        ('⚠️', 'Tarifa espera adicional / 30 min (COP)',''),
        ('💬', '¿Algo especial de este vehículo\nque el sistema deba mencionar al cliente?', ''),
    ]

    t = doc.add_table(rows=len(campos) + 1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezado
    header_texts = ['', 'Campo', 'Respuesta del gerente']
    header_colors = [NEGRO_HEX, NEGRO_HEX, NEGRO_HEX]
    for ci, (ht, hc) in enumerate(zip(header_texts, header_colors)):
        cell = t.cell(0, ci)
        set_cell_bg(cell, hc)
        ph = cell.paragraphs[0]
        ph.paragraph_format.space_before = Pt(6)
        ph.paragraph_format.space_after  = Pt(6)
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_formatted(ph, ht, bold=True, size=9.5,
                          color=BLANCO, font='Calibri')

    # Filas
    for ri, (icono, campo, sugerido) in enumerate(campos, start=1):
        es_abierta = (icono == '💬')
        bg = GRIS_CLARO_HEX if ri % 2 == 0 else BLANCO_HEX

        c0 = t.cell(ri, 0)
        set_cell_bg(c0, bg)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(5)
        p0.paragraph_format.space_after  = Pt(5)
        add_run_formatted(p0, icono, size=10, font='Calibri')

        c1 = t.cell(ri, 1)
        set_cell_bg(c1, bg)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_before = Pt(5)
        p1.paragraph_format.space_after  = Pt(5)
        add_run_formatted(p1, campo, bold=True, size=9.5,
                          color=NEGRO, font='Calibri')

        c2 = t.cell(ri, 2)
        set_cell_bg(c2, AMARILLO_SUAVE if es_abierta else BLANCO_HEX)
        p2 = c2.paragraphs[0]
        p2.paragraph_format.space_before = Pt(5)
        p2.paragraph_format.space_after  = Pt(5)
        if sugerido:
            add_run_formatted(p2, f'Ej: {sugerido}', italic=True, size=9,
                              color=GRIS_TEXT, font='Calibri')
        else:
            add_run_formatted(p2, '✏️  ___________________________',
                              size=9, color=GRIS_TEXT, font='Calibri')

    # Anchos de columna
    for ri2 in range(len(campos) + 1):
        t.cell(ri2, 0).width = Cm(1.0)
        t.cell(ri2, 1).width = Cm(7.5)
        t.cell(ri2, 2).width = Cm(8.5)

    add_horizontal_line(doc, DORADO_HEX, '4')
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def tabla_verificacion(doc, filas):
    """Tabla de verificación para datos existentes."""
    t = doc.add_table(rows=len(filas) + 1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Concepto', 'Valor actual', '¿Correcto?', 'Corrección (si aplica)']
    for ci, h in enumerate(headers):
        cell = t.cell(0, ci)
        set_cell_bg(cell, NEGRO_HEX)
        ph = cell.paragraphs[0]
        ph.paragraph_format.space_before = Pt(6)
        ph.paragraph_format.space_after  = Pt(6)
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_formatted(ph, h, bold=True, size=9, color=BLANCO, font='Calibri')

    for ri, (concepto, valor) in enumerate(filas, start=1):
        bg = GRIS_CLARO_HEX if ri % 2 == 0 else BLANCO_HEX
        datos_fila = [concepto, valor, '  ✅ Sí  /  ❌ No', '✏️  ___________________']
        for ci, texto in enumerate(datos_fila):
            cell = t.cell(ri, ci)
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after  = Pt(5)
            bold_col = ci in (0,)
            add_run_formatted(p, texto, bold=bold_col, size=9,
                              color=NEGRO, font='Calibri')

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def bloque_tipo_servicio(doc, letra, nombre, descripcion, caso_ejemplo, preguntas):
    subtitulo(doc, f'TIPO {letra} — {nombre}', descripcion)
    nota_ejemplo(doc, caso_ejemplo)
    for pregunta, obligatorio, abierta in preguntas:
        campo_respuesta(doc, pregunta, obligatorio=obligatorio, es_abierta=abierta)


def bloque_footer(doc):
    add_horizontal_line(doc, DORADO_HEX, '8')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    add_run_formatted(p, 'DESTINOS EXPRESS S.A.S.  ·  NIT: 900.982.154-2  ·  Bogotá D.C., Colombia',
                      size=8, color=GRIS_TEXT, font='Calibri')
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_formatted(p2,
                      'comercial@destinosexpress.com  ·  +57 302 4060101  ·  www.destinosexpress.com',
                      size=8, color=GRIS_TEXT, font='Calibri')
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_formatted(p3, 'Documento de uso interno — Marzo 2026',
                      italic=True, size=8, color=RGBColor(0xAA, 0xAA, 0xAA), font='Calibri')


# ─── MAIN ─────────────────────────────────────────────────────────────────────

def generar_brief():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── PORTADA ───────────────────────────────────────────────────────────────
    bloque_portada(doc)
    add_page_break(doc)

    # ── SECCIÓN 1 — FLOTA ─────────────────────────────────────────────────────
    titulo_seccion(doc, '1', 'FLOTA DE VEHÍCULOS',
                   'Complete una ficha por cada vehículo que opera la empresa.')

    add_paragraph(doc,
        'A continuación encontrará una ficha por vehículo. '
        'Complete todos los campos con los valores reales que maneja su empresa. '
        'Si tiene más vehículos de los que se muestran, puede duplicar la ficha al final.',
        size=9.5, color=GRIS_TEXT, space_after=12)

    vehiculos = [
        ('Vehículo 1 — Camioneta / SUV',        'Camioneta Ejecutiva / SUV', '1 – 4 pasajeros'),
        ('Vehículo 2 — Van Changan',             'Van Changan Ejecutiva',     '5 – 7 pasajeros'),
        ('Vehículo 3 — Van Traffic',             'Van Traffic',               '7 – 9 pasajeros'),
        ('Vehículo 4 — Van Grande',              '',                          '9 – 12 pasajeros'),
        ('Vehículo 5 — Bus Especial',            'Bus Especial',              '20 – 40 pasajeros'),
        ('Vehículo 6 — (si tiene otro)',         '',                          ''),
    ]

    for titulo_v, nombre_s, cap_s in vehiculos:
        subtitulo(doc, titulo_v)
        tabla_vehiculo(doc, titulo_v, nombre_s, cap_s)

    add_page_break(doc)

    # ── SECCIÓN 2 — TIPOS DE SERVICIO ─────────────────────────────────────────
    titulo_seccion(doc, '2', 'TIPOS DE SERVICIO',
                   'Necesitamos entender cómo se cotizan los distintos tipos de servicio que ofrece.')

    # TIPO A
    bloque_tipo_servicio(
        doc,
        letra='A',
        nombre='Servicio con conductor disponible (pernocta en otra ciudad)',
        descripcion=(
            'Cuando el conductor se desplaza a otra ciudad y permanece allí '
            'hasta la fecha de regreso del cliente.'
        ),
        caso_ejemplo=(
            'Cliente pide Bogotá → Ibagué el 5 de marzo a las 6am. '
            'El conductor se queda en Ibagué. Regreso el 7 de marzo a las 6pm.'
        ),
        preguntas=[
            ('¿Cómo se cotiza este servicio? ¿Qué incluye la tarifa del conductor '
             'mientras está en la otra ciudad? (Ej: hospedaje, alimentación, valor por día de disponibilidad)',
             True, True),
            ('¿Hay un valor fijo por día de disponibilidad del conductor? ¿Cuánto?',
             True, False),
            ('¿Se cobra el trayecto de regreso vacío (sin el cliente)?',
             True, False),
            ('¿Hay un mínimo de días para este tipo de servicio?',
             False, False),
            ('¿Cómo se maneja si el cliente cancela uno de los días o cambia la fecha de regreso?',
             False, True),
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # TIPO B
    bloque_tipo_servicio(
        doc,
        letra='B',
        nombre='Servicio por horas / disposición en ciudad',
        descripcion='El vehículo queda a disposición del cliente por un número de horas en la misma ciudad.',
        caso_ejemplo='Cliente necesita el vehículo por 4 horas en Bogotá para visitas de negocios.',
        preguntas=[
            ('¿Cuál es el mínimo de horas para este tipo de servicio?', True, False),
            ('¿La tarifa por hora varía según el vehículo o es igual para todos?', True, False),
            ('¿Hay un costo diferente para la "hora adicional" después del mínimo?', False, False),
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # TIPO C
    bloque_tipo_servicio(
        doc,
        letra='C',
        nombre='Traslado a aeropuerto',
        descripcion='Servicios específicos de traslado hacia o desde el aeropuerto El Dorado u otros.',
        caso_ejemplo='Cliente de El Nogal necesita traslado a El Dorado a las 4:30am.',
        preguntas=[
            ('¿Manejan tarifas fijas para aeropuerto o se calcula por km como cualquier servicio?',
             True, False),
            ('¿Hay diferencia en la tarifa para vuelos muy temprano (antes de las 5am)?',
             False, False),
            ('¿Se cobra el tiempo de espera si el vuelo se retrasa? ¿Cómo?', False, True),
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # TIPO D
    bloque_tipo_servicio(
        doc,
        letra='D',
        nombre='Otros servicios especiales',
        descripcion='Cualquier otro tipo de servicio que ofrezcan y que tenga una tarifa o lógica diferente.',
        caso_ejemplo='Turismo, bodas, eventos, contratos corporativos mensuales, servicios VIP, etc.',
        preguntas=[
            ('Describa cualquier tipo de servicio adicional que no encaje en los anteriores.',
             False, True),
            ('¿Tienen contratos corporativos con empresas? ¿Cómo se manejan las tarifas en ese caso?',
             False, True),
        ]
    )

    add_page_break(doc)

    # ── SECCIÓN 3 — RECARGOS ──────────────────────────────────────────────────
    titulo_seccion(doc, '3', 'RECARGOS Y CONDICIONES ESPECIALES',
                   'Verifique los recargos actuales y agregue los que falten.')

    add_paragraph(doc,
        'Los siguientes recargos están configurados actualmente en el sistema. '
        'Por favor indique si son correctos o si necesitan ajuste.',
        size=9.5, color=GRIS_TEXT, space_after=10)

    tabla_verificacion(doc, [
        ('Horario nocturno (10pm – 5am)', '+10% sobre el valor base'),
        ('Festivos',                      '+10% sobre el valor base'),
        ('Zona rural / vereda',           '+10% a +20% sobre el valor base'),
    ])

    campo_respuesta(doc,
        '¿Hay algún recargo adicional que no esté en la lista anterior? '
        '(Ej: peajes, espera en aeropuerto, zona de difícil acceso, clima, carga especial)',
        obligatorio=False, es_abierta=True)

    # ── SECCIÓN 4 — RUTAS FRECUENTES ─────────────────────────────────────────
    titulo_seccion(doc, '4', 'RUTAS FRECUENTES',
                   'Valide las rutas frecuentes actuales y agregue las más solicitadas por sus clientes.')

    add_paragraph(doc,
        'Estas rutas están actualmente en el sistema como referencia rápida de precios:',
        size=9.5, color=GRIS_TEXT, space_after=8)

    tabla_verificacion(doc, [
        ('Bogotá → Chía',     'Camioneta: $90.000 – $110.000'),
        ('Bogotá → Tunja',    'Camioneta: $380.000 – $420.000  /  Van: $650.000 – $750.000'),
        ('Bogotá → Girardot', 'Camioneta: $320.000 – $360.000  /  Van: $520.000 – $620.000'),
        ('Bogotá → Paipa',    'Camioneta: $420.000 – $480.000  /  Van: $700.000 – $820.000'),
    ])

    campo_respuesta(doc,
        '¿Qué rutas son las más frecuentes en su operación? '
        '¿Hay rutas que faltan o que tienen precios diferentes a los mostrados?',
        obligatorio=False, es_abierta=True)

    # ── SECCIÓN 5 — DATOS DE LA EMPRESA ───────────────────────────────────────
    titulo_seccion(doc, '5', 'INFORMACIÓN DE LA EMPRESA',
                   'Verifique que los datos de contacto en el sistema sean correctos.')

    tabla_verificacion(doc, [
        ('Nombre',            'DESTINOS EXPRESS S.A.S.'),
        ('NIT',               '900.982.154-2'),
        ('Ciudad',            'Bogotá D.C., Colombia'),
        ('Teléfono WhatsApp', '+57 302 4060101'),
        ('Email comercial',   'comercial@destinosexpress.com'),
        ('Página web',        'www.destinosexpress.com'),
    ])

    campo_respuesta(doc,
        '¿Hay un segundo número de teléfono, WhatsApp adicional o email de servicio al cliente?',
        obligatorio=False, es_abierta=False)

    # ── SECCIÓN 6 — PREGUNTAS FINALES ─────────────────────────────────────────
    titulo_seccion(doc, '6', 'PREGUNTAS FINALES',
                   'Información adicional para ajustar el sistema a su operación real.')

    preguntas_finales = [
        ('¿El sistema debe cotizar servicios fuera de Colombia? '
         '(Ej: Ecuador, Venezuela, Panamá)',
         False, False),
        ('¿Tienen clientes corporativos con tarifas especiales o descuentos negociados? '
         '¿Cómo se manejan?',
         False, True),
        ('¿Hay restricciones de zonas donde NO prestan el servicio?',
         False, False),
        ('¿Qué información adicional considera importante que el sistema conozca '
         'sobre la operación de Destinos Express?',
         False, True),
    ]

    for pregunta, oblig, abierta in preguntas_finales:
        campo_respuesta(doc, pregunta, obligatorio=oblig, es_abierta=abierta)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    bloque_footer(doc)

    # ── GUARDAR ───────────────────────────────────────────────────────────────
    ruta = 'brief_gerente_destinos_express.docx'
    doc.save(ruta)
    print(f'Documento generado: {ruta}')


if __name__ == '__main__':
    generar_brief()
